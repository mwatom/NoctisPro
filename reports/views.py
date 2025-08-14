from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.db.models import Count, Q
from django.utils import timezone
from .models import Report, ReportTemplate
from worklist.models import Study
from accounts.models import User
import io
from django.core.files.base import ContentFile
from django.utils.text import slugify
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None
try:
    from docx import Document
except Exception:
    Document = None

@login_required
def report_list(request):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    """List all reports"""
    # Get filter parameters
    search_query = request.GET.get('search', '')
    status_filter = request.GET.get('status', '')
    modality_filter = request.GET.get('modality', '')
    
    # Base queryset
    reports = Report.objects.select_related('study', 'study__patient', 'study__modality', 'radiologist').all()
    
    # Apply filters
    if search_query:
        reports = reports.filter(
            Q(study__patient__first_name__icontains=search_query) |
            Q(study__patient__last_name__icontains=search_query) |
            Q(study__accession_number__icontains=search_query) |
            Q(radiologist__first_name__icontains=search_query) |
            Q(radiologist__last_name__icontains=search_query)
        )
    
    if status_filter:
        reports = reports.filter(status=status_filter)
    
    if modality_filter:
        reports = reports.filter(study__modality__code=modality_filter)
    
    # Order by most recent
    reports = reports.order_by('-report_date')
    
    # Calculate statistics
    total_reports = Report.objects.count()
    draft_reports = Report.objects.filter(status='draft').count()
    pending_reports = Report.objects.filter(status='preliminary').count()
    final_reports = Report.objects.filter(status='final').count()
    
    context = {
        'reports': reports,
        'total_reports': total_reports,
        'draft_reports': draft_reports,
        'pending_reports': pending_reports,
        'final_reports': final_reports,
        'search_query': search_query,
        'status_filter': status_filter,
        'modality_filter': modality_filter,
    }
    
    return render(request, 'reports/report_list.html', context)

@login_required
def write_report(request, study_id):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    """Write report for study"""
    study = get_object_or_404(Study, id=study_id)
    
    # When opening editor, mark study as in progress for editors
    if study.status in ['scheduled', 'suspended']:
        study.status = 'in_progress'
        study.save(update_fields=['status'])
    
    # Try to get existing report or create new one
    try:
        report = Report.objects.get(study=study)
        is_new_report = False
    except Report.DoesNotExist:
        report = None
        is_new_report = True
    
    if request.method == 'POST':
        # Get form data
        clinical_history = request.POST.get('clinical_history', '')
        technique = request.POST.get('technique', '')
        comparison = request.POST.get('comparison', '')
        findings = request.POST.get('findings', '')
        impression = request.POST.get('impression', '')
        recommendations = request.POST.get('recommendations', '')
        status = request.POST.get('status', 'draft')
        action = request.POST.get('action', 'save')
        
        if is_new_report:
            # Create new report
            report = Report.objects.create(
                study=study,
                radiologist=request.user,
                clinical_history=clinical_history,
                technique=technique,
                comparison=comparison,
                findings=findings,
                impression=impression,
                recommendations=recommendations,
                status=status
            )
            messages.success(request, 'Report created successfully!')
        else:
            # Update existing report
            report.clinical_history = clinical_history
            report.technique = technique
            report.comparison = comparison
            report.findings = findings
            report.impression = impression
            report.recommendations = recommendations
            report.status = status
            report.last_modified = timezone.now()
            
            # If finalizing the report
            if status == 'final' or (action == 'submit' and status == 'final'):
                report.signed_date = timezone.now()
            
            report.save()
            messages.success(request, 'Report updated successfully!')
        
        # Update study status when report finalized
        if status == 'final' or (action == 'submit' and status == 'final'):
            study.status = 'completed'
            study.save(update_fields=['status'])
        
        # Redirect based on action
        if action == 'submit':
            messages.success(request, 'Report submitted successfully!')
            return redirect('reports:report_list')
        else:
            # Stay on the same page for continued editing
            return redirect('reports:write_report', study_id=study_id)
    
    context = {
        'study': study,
        'report': report,
        'is_new_report': is_new_report,
    }
    
    return render(request, 'reports/write_report.html', context)

@login_required
def print_report_stub(request, study_id):
    """Simple printable page for facility users; if report exists include it, else show study details and clinical info."""
    study = get_object_or_404(Study, id=study_id)
    report = Report.objects.filter(study=study).first()
    html = f"""
    <html>
      <head>
        <title>Print Study {study.accession_number}</title>
        <style>
          body {{ font-family: Arial, sans-serif; color: #000; }}
          .header {{ display:flex; justify-content: space-between; border-bottom:1px solid #000; padding-bottom:6px; margin-bottom:10px; }}
          .section {{ margin-bottom: 12px; }}
          .label {{ font-weight:bold; }}
          pre {{ white-space: pre-wrap; font-family: inherit; }}
          @media print {{ .noprint {{ display:none; }} }}
        </style>
      </head>
      <body>
        <div class="header">
          <div>
            <div class="label">Facility:</div>
            <div>{study.facility.name}</div>
          </div>
          <div style="text-align:right">
            <div class="label">Accession:</div>
            <div>{study.accession_number}</div>
          </div>
        </div>
        <div class="section"><span class="label">Patient:</span> {study.patient.full_name} ({study.patient.patient_id})</div>
        <div class="section"><span class="label">Modality:</span> {study.modality.code} &nbsp; <span class="label">Date:</span> {study.study_date}</div>
        <div class="section"><span class="label">Priority:</span> {study.priority.upper()}</div>
        <div class="section"><span class="label">Clinical Information:</span><br/><pre>{(study.clinical_info or '').strip() or '-'}</pre></div>
        {f'<div class="section"><span class="label">Findings:</span><br/><pre>{(report.findings or '').strip()}</pre></div>' if report else ''}
        {f'<div class="section"><span class="label">Impression:</span><br/><pre>{(report.impression or '').strip()}</pre></div>' if report else ''}
        <div class="noprint"><button onclick="window.print()">Print</button></div>
      </body>
    </html>
    """
    return HttpResponse(html)


@login_required
def export_report_pdf(request, study_id):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    study = get_object_or_404(Study, id=study_id)
    report = Report.objects.filter(study=study).first()
    html = f"""
    <h2 style='margin:0'>Radiology Report</h2>
    <div><b>Patient:</b> {study.patient.full_name} ({study.patient.patient_id})</div>
    <div><b>Accession:</b> {study.accession_number} &nbsp; <b>Modality:</b> {study.modality.code} &nbsp; <b>Date:</b> {study.study_date}</div>
    <div><b>Priority:</b> {study.priority.upper()}</div>
    <hr/>
    <div><b>Clinical History</b><br/>{(report.clinical_history if report else (study.clinical_info or '')) or '-'}</div>
    <div><b>Technique</b><br/>{(report.technique if report else '') or '-'}</div>
    <div><b>Comparison</b><br/>{(report.comparison if report else '') or '-'}</div>
    <div><b>Findings</b><br/>{(report.findings if report else '') or '-'}</div>
    <div><b>Impression</b><br/>{(report.impression if report else '') or '-'}</div>
    <div><b>Recommendations</b><br/>{(report.recommendations if report else '') or '-'}</div>
    """
    filename = f"report_{slugify(study.accession_number)}.pdf"
    if fitz is None:
        return JsonResponse({'error': 'PDF export not available (PyMuPDF missing).'}, status=500)
    try:
        doc = fitz.open()
        page = doc.new_page()
        # Simple HTML rendering: use a text writer for robustness
        text = fitz.TextWriter(page.rect)
        y = 36
        for line in html.replace('<br/>', '\n').replace('<hr/>', '\n' + '-'*80 + '\n').replace('<b>','').replace('</b>','').replace('<h2','\n<h2').split('\n'):
            if '<h2' in line:
                line = line.replace("</h2>", '').split('>')[-1]
                text.append((36, y), line, fontsize=16)
                y += 28
            else:
                text.append((36, y), fitz.strip_html(line).strip(), fontsize=11)
                y += 16
            if y > page.rect.height - 36:
                page = doc.new_page(); text = fitz.TextWriter(page.rect); y = 36
        page.insert_text((0,0), "")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(buf.getvalue(), content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="{filename}"'
        return resp
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def export_report_docx(request, study_id):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    if Document is None:
        return JsonResponse({'error': 'DOCX export not available (python-docx missing).'}, status=500)
    study = get_object_or_404(Study, id=study_id)
    report = Report.objects.filter(study=study).first()
    doc = Document()
    doc.add_heading('Radiology Report', 0)
    doc.add_paragraph(f"Patient: {study.patient.full_name} ({study.patient.patient_id})")
    doc.add_paragraph(f"Accession: {study.accession_number}    Modality: {study.modality.code}    Date: {study.study_date}")
    doc.add_paragraph(f"Priority: {study.priority.upper()}")
    doc.add_paragraph('')
    sections = [
        ('Clinical History', (report.clinical_history if report else (study.clinical_info or '')) or '-'),
        ('Technique', (report.technique if report else '') or '-'),
        ('Comparison', (report.comparison if report else '') or '-'),
        ('Findings', (report.findings if report else '') or '-'),
        ('Impression', (report.impression if report else '') or '-'),
        ('Recommendations', (report.recommendations if report else '') or '-'),
    ]
    for title, content in sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    resp = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition'] = f'attachment; filename="report_{slugify(study.accession_number)}.docx"'
    return resp
